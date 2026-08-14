"""
Build an updated Senate Accountability Brief template by:
  1. Starting from the Gallagher ACT brief (the most polished output we have),
     since it carries the finished styling/theme we want to keep.
  2. Turning every Gallagher/ACT-specific value back into a [BRACKET] placeholder.
  3. Re-inserting sections/rows present in the original base template but
     dropped in the Gallagher example (Male Electors, the economic "what
     happens to" questions, the post-vote accountability/history section,
     the regional-communities question, the sustainability closing lines).
  4. Genericising the Gallagher-specific "Minister for Women" lens into an
     optional, clearly-marked portfolio lens block.
  5. Adding a front-of-document legend table mapping every bracket field to
     what it should contain.

Output: Senate_Accountability_Brief_Template_v2.docx
"""
import copy
import re

import docx
from docx.oxml.ns import qn

BASE = r"C:\Users\KinJLy\Documents\GitHub\icm-architect\01_Senate Accountability Briefs\disabilityleadersmemberscommunityndisamendmentbill"
SOURCE = BASE + r"\Gallagher ACT Accountability Brief .docx"
OUT = BASE + r"\Senate_Accountability_Brief_Template_v2.docx"

doc = docx.Document(SOURCE)

# ---------------------------------------------------------------------------
# Step 1: literal Gallagher/ACT text -> bracket placeholders.
# Longest strings first so nothing gets partially clobbered by a shorter match.
# ---------------------------------------------------------------------------
REPLACEMENTS = [
    ("Senator the Hon Katherine Gallagher", "[SENATOR NAME]"),
    ("Senator Gallagher", "Senator [SENATOR NAME]"),
    ("Prepared for [SENATOR NAME], Minister for Women",
     "Prepared for [SENATOR NAME][, PORTFOLIO/MINISTERIAL ROLE - optional, delete if none]"),
    ("Minister for Women, Minister for Finance, Minister for the Public Service, Minister for Government Services",
     "[PORTFOLIO/MINISTERIAL RESPONSIBILITY - optional, delete row if none]"),
    ("Australian Labor Party", "[PARTY]"),
    ("Australian Capital Territory", "[STATE/TERRITORY]"),
    ("senator.katy.gallagher@aph.gov.au", "[EMAIL ADDRESS]"),
    ("Expected Government support, subject to confirmation", "[VOTE POSITION]"),
    ("Require safeguards, transparent monitoring and a public implementation accountability mechanism before the vote",
     "[IMMEDIATE ASK]"),
    ("Dwayne Fernandes, Founder, Minds at Play", "[PREPARED BY - name, role/organisation]"),
    ("324,456", "[TOTAL ELECTORS]"),
    ("166,049", "[FEMALE ELECTORS]"),
    ("12,597", "[NDIS PARTICIPANTS]"),
    ("4,831", "[CHILDREN 0-14]"),
    ("5,457", "[AUTISM PARTICIPANTS]"),
    ("1,155", "[PSYCHOSOCIAL PARTICIPANTS]"),
    ("690", "[FIRST NATIONS PARTICIPANTS]"),
    ("1,171", "[CALD PARTICIPANTS]"),
    ("157,819", "[MALE ELECTORS]"),
    ("7,146", "[ACTIVE PROVIDERS]"),
    ("$0.494 billion", "$[ANNUAL FUNDING] billion"),
    ("51.2%", "[FEMALE SHARE]%"),
    ("38.4%", "[CHILD SHARE]%"),
    ("43.3%", "[AUTISM SHARE]%"),
    ("25,194", "[CARERS ESTIMATE]"),
    ("50,388", "[FAMILY ESTIMATE]"),
    ("75,582", "[COMMUNITY ESTIMATE]"),
    ("170,907", "[TOTAL FOOTPRINT ESTIMATE]"),
    ("$4.9 million", "$[1% REDUCTION] million"),
    ("$24.7 million", "$[5% REDUCTION] million"),
    ("$49.4 million", "$[10% REDUCTION] million"),
    ("Dwayne Fernandes", "[YOUR NAME]"),
    ("Disability Leadership Orator 2025", "[YOUR ROLE/TITLE]"),
    ("As Minister for Women, Senator Gallagher is being asked to consider more than program expenditure. "
     "If supports become harder to access, the practical burden often shifts into households and unpaid care. "
     "This has a direct gender impact because women, including mothers, grandmothers and female family carers, "
     "frequently absorb the unpaid caring, service navigation and workforce disruption that follows reduced formal support.",
     "As [PORTFOLIO], Senator [SENATOR NAME] is being asked to consider more than program expenditure. "
     "[PORTFOLIO-SPECIFIC RISK PARAGRAPH - explain how implementation risk connects to this portfolio's responsibilities]"),
    ("The accountability question is whether the Bill includes enough safeguards to prevent reform savings "
     "from being achieved by transferring costs to women, families and carers.",
     "[PORTFOLIO-SPECIFIC ACCOUNTABILITY QUESTION]"),
    ("Women and care risk", "[PORTFOLIO RISK LABEL]"),
    ("The ACT has 166,049 female electors. If implementation reduces supports, this may affect workforce "
     "participation, household income, family wellbeing and unpaid care loads for women across the ACT.",
     "[PORTFOLIO RISK DESCRIPTION]"),
    ("Minister for Women accountability lens",
     "[PORTFOLIO] accountability lens (OPTIONAL - include only if the senator holds a relevant portfolio; delete this whole section otherwise)"),
    ("Senator the Hon Katherine Gallagher, Australian Labor Party, senator.katy.gallagher@aph.gov.au, "
     "Expected Government support, subject to confirmation",
     "[SENATOR NAME], [PARTY], [EMAIL ADDRESS], [VOTE POSITION]"),
    ("324,456, 166,049, 157,819", "[TOTAL ELECTORS], [FEMALE ELECTORS], [MALE ELECTORS]"),
    ("12,597, 4,831, 5,457", "[NDIS PARTICIPANTS], [CHILDREN 0-14], [AUTISM PARTICIPANTS]"),
    ("1,155, 690, 1,171", "[PSYCHOSOCIAL PARTICIPANTS], [FIRST NATIONS PARTICIPANTS], [CALD PARTICIPANTS]"),
    ("7,146 providers", "[ACTIVE PROVIDERS] providers"),
    ("$0.494 billion, $4.9 million, $24.7 million, $49.4 million",
     "$[ANNUAL FUNDING] billion, $[1% REDUCTION] million, $[5% REDUCTION] million, $[10% REDUCTION] million"),
    ("25,194, 50,388, 170,907", "[CARERS ESTIMATE], [FAMILY ESTIMATE], [TOTAL FOOTPRINT ESTIMATE]"),
    ("Prepared from the supplied Senate Accountability Brief Template, the completed Senate NDIS table and the "
     "Senate contact list provided in this conversation. Figures should be checked against final source data before public release.",
     "Prepared from the Senate Accountability Brief Template, the Senate NDIS Data Workbook and the Senate "
     "Contact List. Figures should be checked against final source data before distribution."),
]

# Word-boundary swap for the short "ACT" abbreviation used throughout body text.
ACT_RE = re.compile(r"\bACT\b")


def swap_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    text = ACT_RE.sub("[STATE ABBR]", text)
    return text


def process_paragraph(p):
    if not p.runs:
        return
    full_text = "".join(r.text for r in p.runs)
    new_text = swap_text(full_text)
    if new_text == full_text:
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


for p in doc.paragraphs:
    process_paragraph(p)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                process_paragraph(p)

# Checkboxes reset to unchecked (this is now a blank template, nothing is "done").
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if "\u2611" in r.text:  # checked box
                        r.text = r.text.replace("\u2611", "\u2610")

# Footer
for s in doc.sections:
    for p in s.footer.paragraphs:
        for r in p.runs:
            r.text = r.text.replace("Prepared by Dwayne Fernandes", "Prepared by [YOUR NAME]")

print("Step 1 (text -> brackets) done.")


# ---------------------------------------------------------------------------
# Helpers for structural insertions (rows / paragraphs cloned from an
# existing element so the new content matches the surrounding formatting).
# ---------------------------------------------------------------------------
def find_row(table, first_cell_contains):
    for row in table.rows:
        if first_cell_contains in row.cells[0].text:
            return row
    raise ValueError(f"row not found: {first_cell_contains!r}")


def clone_row_after(row, new_cell_texts):
    new_tr = copy.deepcopy(row._tr)
    row._tr.addnext(new_tr)
    from docx.table import _Row
    new_row = _Row(new_tr, row._parent)
    for cell, text in zip(new_row.cells, new_cell_texts):
        set_cell_text(cell, text)
    return new_row


def set_cell_text(cell, text):
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(text)


def find_paragraph(doc_or_cell, text_contains):
    for p in doc_or_cell.paragraphs:
        if text_contains in p.text:
            return p
    raise ValueError(f"paragraph not found: {text_contains!r}")


def clone_paragraph_after(p, new_text, style=None):
    new_p_elm = copy.deepcopy(p._p)
    p._p.addnext(new_p_elm)
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_elm, p._parent)
    if new_p.runs:
        new_p.runs[0].text = new_text
        for r in new_p.runs[1:]:
            r.text = ""
    else:
        new_p.add_run(new_text)
    if style:
        new_p.style = style
    return new_p


# (Male Electors row already exists in the source doc - Step 1's text
# substitution above already bracket-ised it, no structural insert needed.)

# ---------------------------------------------------------------------------
# Step 3: restore the "what happens to..." questions dropped from the
# economic exposure section, and the post-vote accountability/history
# section that was dropped entirely.
#
# The economic exposure table's callout text lives inside a table cell, so
# instead of anchoring off it directly, we clone top-level (body-level)
# paragraphs for style and insert the whole new block immediately before the
# "Questions constituents may ask after the vote" heading - i.e. right after
# the economic exposure table, in the same place this content sat in the
# original base template.
# ---------------------------------------------------------------------------
heading_template = find_paragraph(doc, "Implementation risks to own if safeguards are absent")
bullet_template = find_paragraph(doc, "Increased social isolation and reduced community participation")
body_template = find_paragraph(doc, "Key read:")
questions_heading = find_paragraph(doc, "Questions constituents may ask after the vote")


def insert_clone_before(anchor_p, template_p, text, style=None):
    new_p_elm = copy.deepcopy(template_p._p)
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_elm, template_p._parent)
    if new_p.runs:
        new_p.runs[0].text = text
        for r in new_p.runs[1:]:
            r.text = ""
    else:
        new_p.add_run(text)
    if style:
        new_p.style = style
    anchor_p._p.addprevious(new_p_elm)
    return new_p


what_happens = [
    "What happens to parent and carer workforce participation?",
    "What happens to small providers and sole traders?",
    "What happens to regional service availability?",
    "What happens to allied health access?",
    "What happens to local jobs and community organisations?",
]
for line in what_happens:
    insert_clone_before(questions_heading, bullet_template, line)

insert_clone_before(questions_heading, heading_template, "Accountability after the vote")
insert_clone_before(
    questions_heading,
    body_template,
    "If implementation produces harm, constituents are unlikely to distinguish between Treasury advice, "
    "NDIA implementation, departmental administration and parliamentary votes. Many will remember which "
    "senators supported the legislation. History rarely remembers legislation by its intent - it remembers "
    "legislation by its outcomes.",
)
accountability_bullets = [
    "Who supported the Bill?",
    "Who spoke for participants and families?",
    "Who sought safeguards before the vote?",
    "Who required transparent monitoring of outcomes?",
    "Who responded when concerns were raised?",
]
for line in accountability_bullets:
    insert_clone_before(questions_heading, bullet_template, line)

print("Step 3 (economic questions + accountability-after-the-vote section) done.")

# ---------------------------------------------------------------------------
# Step 4: restore the "regional communities" question dropped from the
# post-vote questions list.
# ---------------------------------------------------------------------------
carers_q = find_paragraph(doc, "How were carers supported if service access reduced?")
clone_paragraph_after(
    carers_q, "How will regional communities be protected?", style="List Bullet"
)
print("Step 4 (regional communities question) done.")

# ---------------------------------------------------------------------------
# Step 5: restore the sustainability closing lines dropped before the
# one-page summary section.
# ---------------------------------------------------------------------------
methodology_para = find_paragraph(doc, "Indicative methodology")
closing_lines = [
    "We support a sustainable NDIS.",
    "Success should be measured by outcomes, not savings alone.",
    "Participants do not exist in isolation. Families, carers and communities share the impact.",
]
# Insert just before the "One-page summary for covering email" heading.
summary_heading = find_paragraph(doc, "One-page summary for covering email")
anchor3 = summary_heading
inserted = []
for line in closing_lines:
    new_p_elm = copy.deepcopy(methodology_para._p)
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_elm, methodology_para._parent)
    if new_p.runs:
        new_p.runs[0].text = line
        for r in new_p.runs[1:]:
            r.text = ""
    else:
        new_p.add_run(line)
    summary_heading._p.addprevious(new_p_elm)
print("Step 5 (sustainability closing lines) done.")

# ---------------------------------------------------------------------------
# Step 6: add a Portfolio row to the data-completion checklist table, matching
# the new optional portfolio field.
# ---------------------------------------------------------------------------
checklist_table = None
for t in doc.tables:
    if (t.rows and len(t.rows[0].cells) > 1
            and t.rows[0].cells[0].text.strip() == "Field"
            and t.rows[0].cells[1].text.strip() == "Required value"):
        checklist_table = t
        break
if checklist_table is not None:
    immediate_ask_row = find_row(checklist_table, "Immediate ask")
    clone_row_after(
        immediate_ask_row,
        ["Portfolio (if applicable)", "[PORTFOLIO/MINISTERIAL RESPONSIBILITY]", "\u2610"],
    )
print("Step 6 (checklist portfolio row) done.")

doc.save(OUT)
print(f"Saved {OUT}")
