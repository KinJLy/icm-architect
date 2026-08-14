"""
Stage 2 (format): render a senator's records/{slug}/record.md into a filled
Word brief, using 01_reference/brief-template.docx.

Reads the record's frontmatter (identity/status fields) AND its body's
"- Label: value" bullets under Pulled data / Computed fields (the body is the
human edit surface - a human may have corrected a number there during
validation, so this reads whatever is actually on disk, not a cached copy).

If the record has no relevant portfolio (`portfolio: null` in frontmatter),
the whole optional "accountability lens" section (heading, paragraphs, table)
and the "Portfolio relevance" row are structurally removed from the output -
not just left blank.

Does NOT check `validated: true` before rendering (useful for testing/review),
but does NOT tick the output's checklist boxes unless the record is validated
- so an unvalidated render is visibly unfinished if opened.

Usage (choose exactly one selection mode; names are record slugs):
    python render_brief_docx.py katherine-gallagher pauline-hanson ...
    python render_brief_docx.py --party ALP
    python render_brief_docx.py --state QLD
    python render_brief_docx.py --all
"""
import copy
import re
import sys
from pathlib import Path

import docx
import yaml
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent.parent.parent  # _system/tools -> _system -> workspace root
REF = BASE / "01_reference"
TEMPLATE = REF / "brief-template.docx"
VOTE_RULES = REF / "vote-position-rules.md"
RECORDS_DIR = BASE / "records"
OUT_DIR = BASE / "02_output"

STATE_TO_ABBR_ALREADY = True  # state_abbr comes straight from frontmatter


def load_immediate_ask():
    text = VOTE_RULES.read_text(encoding="utf-8")
    m = re.search(r"## Fixed immediate ask.*?\n>\s*(.+)", text, re.S)
    if not m:
        raise RuntimeError("Could not find fixed immediate ask in vote-position-rules.md")
    return m.group(1).strip()


def parse_record(path: Path):
    text = path.read_text(encoding="utf-8")
    fm_text, body = text.split("---", 2)[1:]
    fm = yaml.safe_load(fm_text)

    bullets = {}
    for line in body.splitlines():
        m = re.match(r"^-\s+([^:]+):\s*(.+)$", line.strip())
        if m:
            bullets[m.group(1).strip()] = m.group(2).strip()

    mech_m = re.search(r"If yes, the specific mechanism \(not a generic paragraph\):\s*(.+)", body)
    mechanism = mech_m.group(1).strip() if mech_m else ""

    return fm, bullets, mechanism


def build_replacements(fm, bullets, mechanism, immediate_ask):
    portfolio = fm.get("portfolio")
    relevant = bool(portfolio)

    r = {
        "[SENATOR NAME]": fm["senator_name"],
        "[PARTY]": fm["party"],
        "[STATE/TERRITORY]": fm["state"],
        "[STATE ABBR]": fm["state_abbr"],
        "[EMAIL ADDRESS]": fm["email"],
        "[VOTE POSITION]": fm["vote_position"],
        "[IMMEDIATE ASK]": immediate_ask,
        "[PREPARED BY - name, role/organisation]": "[PREPARED BY - name, role/organisation]",
        "[TOTAL ELECTORS]": bullets["Total Electors"],
        "[FEMALE ELECTORS]": bullets["Female Electors"],
        "[MALE ELECTORS]": bullets["Male Electors"],
        "[NDIS PARTICIPANTS]": bullets["Active NDIS Participants"],
        "[CHILDREN 0-14]": bullets["Children (0-14 Years)"],
        "[AUTISM PARTICIPANTS]": bullets["Autism Participants"],
        "[PSYCHOSOCIAL PARTICIPANTS]": bullets["Psychosocial Participants"],
        "[FIRST NATIONS PARTICIPANTS]": bullets["First Nations Participants"],
        "[CALD PARTICIPANTS]": bullets["CALD Participants"],
        "[ACTIVE PROVIDERS]": bullets["Active Providers"],
        "[ANNUAL FUNDING]": bullets["Annual NDIS Funding Exposure ($ billion)"],
        "[FEMALE SHARE]": bullets["Female Share (%)"],
        "[CHILD SHARE]": bullets["Child Share (%)"],
        "[AUTISM SHARE]": bullets["Autism Share (%)"],
        "[1% REDUCTION]": bullets["1% Reduction ($ million)"],
        "[5% REDUCTION]": bullets["5% Reduction ($ million)"],
        "[10% REDUCTION]": bullets["10% Reduction ($ million)"],
        "[CARERS ESTIMATE]": bullets["Carers Estimate"],
        "[FAMILY ESTIMATE]": bullets["Family Estimate"],
        "[COMMUNITY ESTIMATE]": bullets["Community Estimate"],
        "[TOTAL FOOTPRINT ESTIMATE]": bullets["Total Footprint Estimate"],
    }
    if relevant:
        r["[PORTFOLIO/MINISTERIAL ROLE]"] = portfolio
        r["[PORTFOLIO/MINISTERIAL RESPONSIBILITY - optional, delete row if none]"] = portfolio
        r["[PORTFOLIO/MINISTERIAL RESPONSIBILITY]"] = portfolio
        r["[PORTFOLIO]"] = portfolio
        r["[PORTFOLIO-SPECIFIC RISK PARAGRAPH — explain how implementation risk connects to this portfolio's responsibilities]"] = mechanism
        r["[PORTFOLIO-SPECIFIC ACCOUNTABILITY QUESTION]"] = mechanism
        r["[PORTFOLIO RISK LABEL]"] = f"{portfolio} accountability risk"
        r["[PORTFOLIO RISK DESCRIPTION]"] = mechanism
    else:
        # snapshot-table row and lens section are removed structurally, but the
        # checklist table's summary row survives either way and needs a value.
        r["[PORTFOLIO/MINISTERIAL RESPONSIBILITY]"] = "N/A"
    return r, relevant


def replace_in_paragraph(paragraph, replacements):
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    if "[" not in full_text:
        return
    new_text = full_text
    for k, v in replacements.items():
        new_text = new_text.replace(k, v)
    if new_text == full_text:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def apply_replacements(doc, replacements):
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
    # Headers/footers are a separate part of the document, not covered by
    # doc.paragraphs/doc.tables above - easy to forget, so handled explicitly.
    for s in doc.sections:
        for p in s.header.paragraphs:
            replace_in_paragraph(p, replacements)
        for p in s.footer.paragraphs:
            replace_in_paragraph(p, replacements)


def strip_prepared_for_clause(doc, relevant, portfolio):
    marker = "[, PORTFOLIO/MINISTERIAL ROLE - optional, delete if none]"
    replacement = f", {portfolio}" if relevant else ""
    for p in doc.paragraphs:
        if marker in "".join(r.text for r in p.runs):
            full = "".join(r.text for r in p.runs)
            p.runs[0].text = full.replace(marker, replacement)
            for r in p.runs[1:]:
                r.text = ""


def strip_lens_optional_note(doc):
    marker = " (OPTIONAL - include only if the senator holds a relevant portfolio; delete this whole section otherwise)"
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if marker in full:
            p.runs[0].text = full.replace(marker, "")
            for r in p.runs[1:]:
                r.text = ""


def remove_portfolio_row(doc):
    for t in doc.tables:
        for row in t.rows:
            if row.cells[0].text.strip() == "Portfolio relevance":
                row._tr.getparent().remove(row._tr)
                return


def remove_portfolio_lens_section(doc):
    body = doc.element.body
    children = list(body.iterchildren())

    def text_of(elm):
        if elm.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            return Paragraph(elm, doc).text
        return None

    start_idx = None
    end_idx = None
    for i, elm in enumerate(children):
        t = text_of(elm)
        if t and "accountability lens" in t and start_idx is None:
            start_idx = i
            continue
        if start_idx is not None and t == "Community footprint":
            end_idx = i
            break
    if start_idx is None or end_idx is None:
        print("WARN: could not locate portfolio-lens section bounds; leaving as-is")
        return
    for elm in children[start_idx:end_idx]:
        body.remove(elm)


def set_checklist_checked(doc, checked: bool):
    if not checked:
        return
    for t in doc.tables:
        if t.rows and t.rows[0].cells[0].text.strip() == "Field" and len(t.rows[0].cells) > 1 and t.rows[0].cells[1].text.strip() == "Required value":
            for row in t.rows[1:]:
                for p in row.cells[-1].paragraphs:
                    for r in p.runs:
                        r.text = r.text.replace("\u2610", "\u2611")


def render(slug: str, immediate_ask: str):
    record_path = RECORDS_DIR / slug / "record.md"
    if not record_path.exists():
        print(f"SKIP (no record): {slug}")
        return
    fm, bullets, mechanism = parse_record(record_path)
    replacements, relevant = build_replacements(fm, bullets, mechanism, immediate_ask)

    doc = docx.Document(str(TEMPLATE))
    strip_prepared_for_clause(doc, relevant, fm.get("portfolio"))
    if not relevant:
        remove_portfolio_row(doc)
        remove_portfolio_lens_section(doc)
    else:
        strip_lens_optional_note(doc)
    apply_replacements(doc, replacements)
    set_checklist_checked(doc, fm.get("validated", False))

    OUT_DIR.mkdir(exist_ok=True)
    out_path = OUT_DIR / f"{fm['senator_name']} Senate Impact Pack.docx"
    doc.save(str(out_path))
    status = "validated" if fm.get("validated") else "NOT YET VALIDATED"
    print(f"Wrote {out_path}  [{status}]")


def scan_existing_records():
    """{slug: {"party": party_abbr, "state": state_abbr}} for every record on disk."""
    available = {}
    for p in RECORDS_DIR.glob("*/record.md"):
        fm, _, _ = parse_record(p)
        available[p.parent.name] = {"party": fm.get("party_abbr", ""), "state": fm.get("state_abbr", "")}
    return available


if __name__ == "__main__":
    sys.path.insert(0, str(Path(__file__).parent))
    from _selection import parse_selection_args

    immediate_ask = load_immediate_ask()
    available = scan_existing_records()
    default_slugs = ["katherine-gallagher", "pauline-hanson", "david-pocock", "jordon-steele-john"]
    selection, defaults = parse_selection_args(sys.argv[1:], default_names=default_slugs)
    slugs = selection.resolve(available, default_keys=defaults)
    print(f"Selected {len(slugs)} record(s)")

    for slug in slugs:
        render(slug, immediate_ask)
